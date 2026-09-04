"""
Módulo 7: GENERACIÓN DE REPORTES TÉCNICOS Y EJECUTIVOS.
- Formatos completamente funcionales:
  1. Excel (.xlsx) con openpyxl y formato institucional.
  2. PDF Ejecutivo (.pdf) con reportlab en orientación horizontal y tablas estilizadas.
  3. Word Editable (.docx) con python-docx y tablas con sombreado UNT.
- Tipos de Reportes:
  1. Estado General de Flota y Activos
  2. Modelos de IA y Evaluación CRISP-DM
  3. Historial de Fallas y Reparaciones
  4. Indicadores Operacionales y KPIs Mineros
  5. Alertas y Eventos Críticos
- Trazabilidad y registro en tabla 'reporte_generado' de PostgreSQL y en memoria.
"""

import io
from datetime import datetime
import streamlit as st
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from config.settings import REPORTS_PDF_DIR, REPORTS_WORD_DIR, REPORTS_EXCEL_DIR
from database.queries import (
    get_equipos, get_alertas, _MEMORY_DB, registrar_reporte_generado
)
from auth.authentication import AuthManager

COLUMNAS_REPORTES = {
    "Estado General de Flota y Activos": [
        "codigo_equipo", "nombre_equipo", "tipo_equipo", "marca",
        "modelo", "horas_operacion", "estado_operativo", "ubicacion_actual"
    ],
    "Modelos de IA y Evaluación CRISP-DM": [
        "Modelo", "Fase CRISP-DM", "Exactitud (Accuracy)", "F1-Score", "ROC-AUC", "Estado"
    ],
    "Historial de Fallas y Reparaciones": [
        "id_falla", "id_equipo", "tipo_falla", "severidad",
        "descripcion_falla", "tiempo_inactividad_minutos", "costo_reparacion", "estado_falla"
    ],
    "Indicadores Operacionales y KPIs Mineros": [
        "id_equipo", "periodo", "disponibilidad", "utilizacion",
        "mtbf_horas", "mttr_horas", "oee", "costo_mantenimiento"
    ],
    "Alertas y Eventos Críticos": [
        "id_alerta", "id_equipo", "tipo_alerta", "nivel_gravedad",
        "mensaje_alerta", "estado_alerta", "fecha_generacion"
    ]
}

def safe_report_dataframe(raw_data, target_columns: list) -> pd.DataFrame:
    """
    Construye un DataFrame asegurando que nunca lance KeyError, 
    incluso si raw_data es None, una lista vacía [], o faltan columnas esperadas.
    """
    if isinstance(raw_data, pd.DataFrame):
        df = raw_data.copy()
    elif not raw_data:
        return pd.DataFrame(columns=target_columns)
    else:
        df = pd.DataFrame(raw_data)

    for col in target_columns:
        if col not in df.columns:
            df[col] = "N/A"

    return df[target_columns]

def generate_excel_report(report_type: str, data: pd.DataFrame) -> bytes:
    """Genera un archivo Excel profesional con openpyxl con estilos corporativos UNT."""
    output = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = report_type[:30]

    # Estilos
    header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    title_font = Font(name="Calibri", size=14, bold=True, color="1E3A8A")
    border_thin = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    # Título institucional
    ws.merge_cells("A1:G1")
    ws["A1"] = "UNIVERSIDAD NACIONAL DE TRUJILLO - SISTEMA DE MANTENIMIENTO PREDICTIVO"
    ws["A1"].font = title_font

    ws.merge_cells("A2:G2")
    ws["A2"] = f"Reporte Técnico: {report_type} | Fecha Emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws["A2"].font = Font(name="Calibri", size=10, italic=True, color="475569")

    # Encabezados de tabla
    start_row = 4
    for col_idx, col_name in enumerate(data.columns, start=1):
        cell = ws.cell(row=start_row, column=col_idx, value=str(col_name).upper())
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Filas de datos
    if data.empty:
        ws.cell(row=start_row + 1, column=1, value="Sin registros registrados en el sistema.")
        ws.cell(row=start_row + 1, column=1).font = Font(name="Calibri", size=10, italic=True, color="64748B")
    else:
        for r_idx, row in enumerate(data.itertuples(index=False), start=start_row + 1):
            for c_idx, val in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=val if pd.notnull(val) else "")
                cell.border = border_thin
                cell.alignment = Alignment(vertical="center")

    # Auto-ajuste de columnas
    for col in ws.columns:
        vals = [str(cell.value or '') for cell in col]
        max_len = max([len(v) for v in vals], default=10) if vals else 10
        col_letter = openpyxl.utils.get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 14)

    wb.save(output)
    return output.getvalue()

def generate_pdf_report(report_type: str, data: pd.DataFrame) -> bytes:
    """Genera un archivo PDF ejecutivo con ReportLab con diseño institucional UNT."""
    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(letter),
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'PdfTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=15,
        textColor=colors.HexColor('#1E3A8A'),
        alignment=1,
        spaceAfter=3
    )
    sub_style = ParagraphStyle(
        'PdfSub',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        textColor=colors.HexColor('#475569'),
        alignment=1,
        spaceAfter=10
    )
    meta_style = ParagraphStyle(
        'PdfMeta',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8.5,
        textColor=colors.HexColor('#64748B'),
        alignment=0,
        spaceAfter=8
    )
    cell_h_style = ParagraphStyle(
        'PdfCellHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        textColor=colors.white,
        alignment=1
    )
    cell_d_style = ParagraphStyle(
        'PdfCellData',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        textColor=colors.HexColor('#0F172A'),
        alignment=0
    )

    elements = [
        Paragraph("UNIVERSIDAD NACIONAL DE TRUJILLO", title_style),
        Paragraph("SISTEMA DE MANTENIMIENTO PREDICTIVO CON INTELIGENCIA ARTIFICIAL (UNT IS-402)", sub_style),
        Paragraph(
            f"<b>Tipo de Reporte:</b> {report_type} &nbsp;&nbsp;|&nbsp;&nbsp; "
            f"<b>Fecha de Emisión:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')} &nbsp;&nbsp;|&nbsp;&nbsp; "
            f"<b>Total Registros:</b> {len(data)}",
            meta_style
        ),
        Spacer(1, 6)
    ]

    # Calcular ancho dinámico por columna (ancho útil: 720 puntos)
    num_cols = max(len(data.columns), 1)
    col_width = 720.0 / num_cols

    table_data = [[Paragraph(f"<b>{str(c).upper()}</b>", cell_h_style) for c in data.columns]]
    if data.empty:
        table_data.append([
            Paragraph("<i>Sin registros</i>", cell_d_style)
            for _ in data.columns
        ])
    else:
        for row in data.itertuples(index=False):
            table_data.append([
                Paragraph(str(val if pd.notnull(val) else "-"), cell_d_style)
                for val in row
            ])

    t = Table(table_data, colWidths=[col_width] * num_cols)
    t_style = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 5),
        ('TOPPADDING', (0, 0), (-1, 0), 5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]
    for i in range(1, len(table_data)):
        bg = colors.HexColor('#F8FAFC') if i % 2 == 0 else colors.white
        t_style.append(('BACKGROUND', (0, i), (-1, i), bg))

    t.setStyle(TableStyle(t_style))
    elements.append(t)

    doc.build(elements)
    return output.getvalue()

def generate_word_report(report_type: str, data: pd.DataFrame) -> bytes:
    """Genera un archivo Word (.docx) editable con python-docx y formato institucional."""
    doc = docx.Document()

    # Orientación horizontal para tablas con más de 4 columnas
    if len(data.columns) > 4:
        section = doc.sections[0]
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    # Título principal institucional
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("UNIVERSIDAD NACIONAL DE TRUJILLO")
    r_title.font.name = 'Calibri'
    r_title.font.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # Subtítulo institucional
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Sistema de Mantenimiento Predictivo con IA - Ingeniería de Software II (IS-402)")
    r_sub.font.name = 'Calibri'
    r_sub.font.italic = True
    r_sub.font.size = Pt(10.5)
    r_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Metadatos del informe
    p_meta = doc.add_paragraph()
    r_m1 = p_meta.add_run("Reporte Técnico: ")
    r_m1.bold = True
    p_meta.add_run(f"{report_type}    |    ")
    r_m2 = p_meta.add_run("Fecha de Emisión: ")
    r_m2.bold = True
    p_meta.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}    |    ")
    r_m3 = p_meta.add_run("Total Registros: ")
    r_m3.bold = True
    p_meta.add_run(str(len(data)))

    # Tabla de datos
    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados con sombreado azul institucional
    for i, col in enumerate(data.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col).upper()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A8A"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Filas de datos con filas alternadas
    if data.empty:
        r_cells = table.add_row().cells
        for c_idx in range(len(data.columns)):
            r_cells[c_idx].text = "Sin registros"
    else:
        for r_idx, row in enumerate(data.itertuples(index=False)):
            r_cells = table.add_row().cells
            shd_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(row):
                cell = r_cells[c_idx]
                cell.text = str(val if pd.notnull(val) else "-")
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shd_color}"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def render_reportes():
    st.markdown("## 📑 Generador de Reportes Técnicos y Ejecutivos")
    st.caption("Exportación multi-formato (Excel .xlsx, PDF .pdf, Word .docx) con trazabilidad en base de datos.")

    user = st.session_state.get("user", {})
    rol = user.get("rol", "Técnico")
    if not AuthManager.check_permission(rol, "generar_reportes"):
        st.warning("Su rol no cuenta con permisos para generar reportes ejecutivos.")
        return

    col_rep, col_fmt = st.columns([3, 2])
    with col_rep:
        tipo_reporte = st.selectbox(
            "Seleccionar Tipo de Reporte:",
            [
                "Estado General de Flota y Activos",
                "Modelos de IA y Evaluación CRISP-DM",
                "Historial de Fallas y Reparaciones",
                "Indicadores Operacionales y KPIs Mineros",
                "Alertas y Eventos Críticos"
            ]
        )
    with col_fmt:
        formato = st.selectbox(
            "Formato de Exportación:",
            ["Excel (.xlsx)", "PDF Ejecutivo (.pdf)", "Word Editable (.docx)"]
        )

    st.markdown("---")

    # Generación y previsualización de datos
    if tipo_reporte == "Estado General de Flota y Activos":
        equipos = get_equipos(incluir_desactivados=True)
        cols = COLUMNAS_REPORTES["Estado General de Flota y Activos"]
        df_export = safe_report_dataframe(equipos, cols)
    elif tipo_reporte == "Modelos de IA y Evaluación CRISP-DM":
        from modules.ia_engine import PredictiveMaintenanceAIEngine
        engine = PredictiveMaintenanceAIEngine()
        df_models = engine.get_models_comparison_table()
        cols = COLUMNAS_REPORTES["Modelos de IA y Evaluación CRISP-DM"]
        if set(cols).issubset(df_models.columns):
            df_export = safe_report_dataframe(df_models, cols)
        else:
            df_export = df_models
    elif tipo_reporte == "Historial de Fallas y Reparaciones":
        cols = COLUMNAS_REPORTES["Historial de Fallas y Reparaciones"]
        df_export = safe_report_dataframe(_MEMORY_DB.get("falla", []), cols)
    elif tipo_reporte == "Indicadores Operacionales y KPIs Mineros":
        cols = COLUMNAS_REPORTES["Indicadores Operacionales y KPIs Mineros"]
        df_export = safe_report_dataframe(_MEMORY_DB.get("kpi_equipo", []), cols)
    else:
        cols = COLUMNAS_REPORTES["Alertas y Eventos Críticos"]
        df_export = safe_report_dataframe(get_alertas(), cols)

    st.markdown("#### 👁️ Vista Previa del Conjunto de Datos")
    if df_export.empty:
        st.info("ℹ️ Actualmente no existen registros en el sistema para este reporte. Se generará la estructura oficial con encabezados correspondientes.")
    st.dataframe(df_export.head(15), use_container_width=True, hide_index=True)

    c_gen1, c_gen2 = st.columns([2, 4])
    with c_gen1:
        if st.button("🚀 Generar y Registrar Reporte", type="primary", use_container_width=True):
            # Determinar extensión, tipo MIME e ícono según formato seleccionado
            if "PDF" in formato:
                ext = "pdf"
                mime = "application/pdf"
                raw_bytes = generate_pdf_report(tipo_reporte, df_export)
                icon = "📕"
            elif "Word" in formato:
                ext = "docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                raw_bytes = generate_word_report(tipo_reporte, df_export)
                icon = "📘"
            else:
                ext = "xlsx"
                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                raw_bytes = generate_excel_report(tipo_reporte, df_export)
                icon = "📗"

            safe_name = tipo_reporte.replace(' ', '_').replace('/', '_')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')
            filename = f"Reporte_{safe_name}_{timestamp}.{ext}"

            # Guardar copia en el directorio correspondiente
            if ext == "pdf":
                file_path = REPORTS_PDF_DIR / filename
            elif ext == "docx":
                file_path = REPORTS_WORD_DIR / filename
            else:
                file_path = REPORTS_EXCEL_DIR / filename

            try:
                with open(file_path, "wb") as f_out:
                    f_out.write(raw_bytes)
            except Exception:
                pass

            # Registrar en bitácora y tabla reporte_generado
            rep_id = registrar_reporte_generado({
                "nombre_reporte": f"{tipo_reporte} - UNT IS402",
                "tipo_reporte": tipo_reporte,
                "formato": ext,
                "id_usuario_genero": user.get("id_usuario", 1),
                "parametros_usados": {"filtro": "todos", "registros": len(df_export)},
                "ruta_archivo": str(file_path),
                "tamano_bytes": len(raw_bytes)
            })

            # Guardar en session_state para descarga inmediata
            st.session_state["reporte_descarga"] = {
                "id": rep_id,
                "bytes": raw_bytes,
                "filename": filename,
                "mime": mime,
                "formato": formato,
                "tipo": tipo_reporte,
                "icon": icon
            }
            st.success(f"¡Reporte {icon} **{filename}** generado exitosamente con ID #{rep_id}!")

    if "reporte_descarga" in st.session_state:
        rep_info = st.session_state["reporte_descarga"]
        with c_gen2:
            st.download_button(
                label=f"📥 Descargar Archivo {rep_info['icon']} ({rep_info['filename']})",
                data=rep_info["bytes"],
                file_name=rep_info["filename"],
                mime=rep_info["mime"],
                type="primary",
                use_container_width=True
            )

    st.markdown("---")
    st.markdown("#### 🗄️ Historial de Reportes Generados en el Sistema (reporte_generado)")
    reps_log = _MEMORY_DB.get("reporte_generado", [])
    if reps_log:
        df_reps = pd.DataFrame(reps_log)
        cols_reps = ["id_reporte", "nombre_reporte", "tipo_reporte", "formato", "fecha_generacion", "veces_descargado", "tamano_bytes"]
        cols_avail = [c for c in cols_reps if c in df_reps.columns]
        st.dataframe(df_reps[cols_avail].sort_values("id_reporte", ascending=False), use_container_width=True, hide_index=True)
    else:
        st.info("No se han emitido reportes en la sesión actual.")
